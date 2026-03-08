"""Fixed-template report export helpers."""

from __future__ import annotations

import json
import re
from datetime import datetime
from pathlib import Path
from typing import Any


FILENAME_PATTERN = re.compile(
    r"^evidence_frame_(?P<frame>\d+)_Track_(?P<track>\d+)_(?P<phase>Pre1|Start|Post1)\.(jpg|jpeg|png)$",
    re.IGNORECASE,
)


def _parse_evidence_name(path: Path) -> tuple[int, int, str] | None:
    match = FILENAME_PATTERN.match(path.name)
    if not match:
        return None
    return int(match.group("frame")), int(match.group("track")), match.group("phase")


def collect_local_evidence_pairs(evidence: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Collect grouped local evidence image pairs from tool outputs."""

    grouped: dict[tuple[int, int], dict[str, Any]] = {}

    for item in evidence:
        if item.get("tool") != "car_detection":
            continue
        detail = item.get("detail") or {}
        for raw_path in detail.get("evidence_images") or []:
            path = Path(str(raw_path))
            parsed = _parse_evidence_name(path)
            if not parsed:
                continue

            frame_index, track_id, phase = parsed
            start_frame = frame_index + 1 if phase == "Pre1" else frame_index
            key = (track_id, start_frame)
            entry = grouped.setdefault(
                key,
                {
                    "track_id": track_id,
                    "start_frame": start_frame,
                    "pre_image_path": None,
                    "start_image_path": None,
                    "pre_image_name": None,
                    "start_image_name": None,
                },
            )

            if phase == "Pre1":
                entry["pre_image_path"] = str(path)
                entry["pre_image_name"] = path.name
            elif phase == "Start":
                entry["start_image_path"] = str(path)
                entry["start_image_name"] = path.name

    pairs = [
        item
        for item in grouped.values()
        if item.get("pre_image_path") or item.get("start_image_path")
    ]
    pairs.sort(key=lambda item: (item.get("track_id", 0), item.get("start_frame", 0)))
    return pairs


def _format_value(value: Any) -> str:
    if value is None:
        return "-"
    if isinstance(value, float):
        return f"{value:.3f}"
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False, indent=2)
    return str(value)


def _tool_summary_lines(evidence: list[dict[str, Any]]) -> list[str]:
    lines: list[str] = []
    for item in evidence:
        tool_name = str(item.get("tool") or "unknown")
        status = str(item.get("status") or "ok")
        detail = item.get("detail")
        lines.append(f"工具：{tool_name}（状态：{status}）")

        if isinstance(detail, dict):
            for key, value in detail.items():
                lines.append(f"- {key}: {_format_value(value)}")
        else:
            lines.append(f"- detail: {_format_value(detail)}")

        lines.append("")
    return lines


def _configure_document_styles(document) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    for style_name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(10.5 if style_name == "Normal" else 12)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _build_docx_report(
    *,
    output_path: Path,
    request_id: str,
    video_locator: str,
    detection_status: dict[str, Any],
    planning_text: str | None,
    initial_tools: list[str],
    deferred_tools: list[str],
    reflection_notes: list[str],
    reflection_decision: str | None,
    evidence: list[dict[str, Any]],
    audit_trail: list[str],
) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    document = Document()
    _configure_document_styles(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("视频篡改检测报告")
    title_run.bold = True
    title_run.font.size = document.styles["Title"].font.size

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    document.add_heading("1. 任务信息", level=1)
    info_table = document.add_table(rows=0, cols=2)
    info_table.style = "Table Grid"
    for label, value in (
        ("请求 ID", request_id),
        ("视频路径", video_locator),
        ("检测结论", detection_status.get("verdict", "unknown")),
        ("置信度", f"{float(detection_status.get('confidence', 0.0)):.1%}"),
    ):
        row = info_table.add_row().cells
        row[0].text = label
        row[1].text = _format_value(value)

    document.add_heading("2. 综合结论", level=1)
    document.add_paragraph(_format_value(detection_status.get("rationale") or "暂无综合说明。"))

    document.add_heading("3. 分析流程", level=1)
    document.add_paragraph(f"初筛工具：{', '.join(initial_tools) if initial_tools else '无'}")
    document.add_paragraph(f"深查工具：{', '.join(deferred_tools) if deferred_tools else '无'}")
    document.add_paragraph(f"反思决策：{reflection_decision or '无'}")
    if reflection_notes:
        for note in reflection_notes:
            document.add_paragraph(note, style="List Bullet")
    if planning_text:
        document.add_paragraph("规划文本：")
        document.add_paragraph(planning_text)

    document.add_heading("4. 工具输出摘要", level=1)
    summary_lines = _tool_summary_lines(evidence)
    if summary_lines:
        for line in summary_lines:
            document.add_paragraph(line)
    else:
        document.add_paragraph("本次任务未返回工具输出。")

    document.add_heading("5. 证据图片", level=1)
    image_pairs = collect_local_evidence_pairs(evidence)
    if not image_pairs:
        document.add_paragraph("未提取到证据图片。")
    else:
        for pair in image_pairs:
            track_id = pair.get("track_id")
            start_frame = pair.get("start_frame")
            document.add_heading(f"Track {track_id} / 起始帧 {start_frame}", level=2)

            for label, key in (("前一帧证据", "pre_image_path"), ("起始帧证据", "start_image_path")):
                image_path_str = pair.get(key)
                if not image_path_str:
                    continue
                image_path = Path(str(image_path_str))
                if not image_path.exists():
                    continue
                document.add_paragraph(f"{label}：{image_path.name}")
                try:
                    document.add_picture(str(image_path), width=Inches(5.8))
                except Exception as exc:
                    document.add_paragraph(f"图片插入失败：{exc}")

    document.add_heading("6. 审计日志", level=1)
    if audit_trail:
        for item in audit_trail:
            document.add_paragraph(str(item), style="List Bullet")
    else:
        document.add_paragraph("无审计日志。")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def _convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> str | None:
    try:
        from docx2pdf import convert
    except ImportError:
        return "未安装 docx2pdf，已生成 Word 报告但未导出 PDF。"

    try:
        convert(str(docx_path), str(pdf_path))
    except Exception as exc:
        return f"PDF 导出失败：{exc}"

    if not pdf_path.exists():
        return "PDF 导出未生成文件，请检查本机 Word / Office 环境。"
    return None


def generate_detection_report(
    *,
    cache_root: Path,
    request_id: str,
    video_locator: str,
    detection_status: dict[str, Any],
    planning_text: str | None,
    initial_tools: list[str],
    deferred_tools: list[str],
    reflection_notes: list[str],
    reflection_decision: str | None,
    evidence: list[dict[str, Any]],
    audit_trail: list[str],
    target_directory: Path | None = None,
) -> dict[str, str | None]:
    """Generate fixed-template report artifacts for one detection request."""

    report_dir = target_directory or (cache_root / "reports" / request_id)
    docx_path = report_dir / "video_tampering_report.docx"
    pdf_path = report_dir / "video_tampering_report.pdf"

    _build_docx_report(
        output_path=docx_path,
        request_id=request_id,
        video_locator=video_locator,
        detection_status=detection_status,
        planning_text=planning_text,
        initial_tools=initial_tools,
        deferred_tools=deferred_tools,
        reflection_notes=reflection_notes,
        reflection_decision=reflection_decision,
        evidence=evidence,
        audit_trail=audit_trail,
    )

    pdf_error = _convert_docx_to_pdf(docx_path, pdf_path)

    return {
        "report_directory": str(report_dir),
        "docx_path": str(docx_path),
        "pdf_path": str(pdf_path) if pdf_path.exists() else None,
        "generation_error": pdf_error,
    }